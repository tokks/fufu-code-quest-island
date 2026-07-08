import os
from docx import Document

DOWNLOADS_DIR = r"C:\Users\97541\Downloads"

for filename in os.listdir(DOWNLOADS_DIR):
    if "day_01" in filename.lower() and "习题答案" not in filename:
        print(f"文件名: {filename}")
        doc = Document(os.path.join(DOWNLOADS_DIR, filename))
        print("\n文档内容:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"{i}: {text}")

print("\n" + "="*50 + "\n")

for filename in os.listdir(DOWNLOADS_DIR):
    if "day_01" in filename.lower() and "习题答案" in filename:
        print(f"答案文件名: {filename}")
        doc = Document(os.path.join(DOWNLOADS_DIR, filename))
        print("\n答案内容:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"{i}: {text}")